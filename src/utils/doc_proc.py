#Document processing utilities.
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import os

def process_document(document_path):
    try:
        _, file_extension = os.path.splitext(document_path)
        if file_extension.lower() in ['.pdf']:
            return process_pdf(document_path)
        elif file_extension.lower() in ['.doc', '.docx']:
            return process_docx(document_path)
        elif file_extension.lower() in ['.jpg', '.jpeg', '.png', '.heif']:
            return process_image(document_path)
        else:
            return {"status": "error", "message": "Unsupported file type"}

    except Exception as e:
        raise ValueError(f"Error in processing document: {str(e)}")

def process_pdf(document_path):
    try:
        with open(document_path, 'rb') as file:
            reader = PdfReader(file)
            text = ''
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text
        if not text.strip():
            raise ValueError("No text found in PDF")
        return {"document_data": extract_titles_and_content(text)}
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")

def process_docx(document_path):
    try:
        doc = DocxDocument(document_path)
        text = ''
        for para in doc.paragraphs:
            text += para.text + '\n'
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\n'

        if not text.strip():
            raise ValueError("No text found in DOCX")
        return {"document_data": extract_titles_and_content(text)}
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")

def process_image(document_path):
    try:
        img = Image.open(document_path)
        text = pytesseract.image_to_string(img)
        if not text.strip():
            raise ValueError("No text found in Image")
        return {"document_data": extract_titles_and_content(text)}
    except Exception as e:
        raise ValueError(f"Error processing image: {str(e)}")

def extract_titles_and_content(text):
    lines = text.split('\n')
    data = {}
    current_title = None

    for line in lines:
        if line.strip() == '':
            continue

        if line.strip().isupper():
            current_title = line.strip()
            data[current_title] = []
        elif current_title:
            data[current_title].append(line.strip())

    return data
